from docx import Document

# Create a new Word document
doc = Document()

# Add name and contact info
doc.add_heading('Himanshu (Heath) Thapa, M.Eng', level=1)
doc.add_paragraph('(+) 1 647-425-7751 | heath.thapa@gmail.com | LinkedIn | Portfolio')

# Profile
doc.add_heading('PROFILE', level=2)
doc.add_paragraph(
    "Marketing Science & Analytics leader with 8+ years of experience measuring and optimizing cross-channel "
    "marketing activations across paid media, CRM, and product engagement. Proven expertise in attribution modeling "
    "(MMM & MTA), campaign ROI analysis, segmentation, experimentation, and marketing performance measurement. "
    "Strong record of integrating digital media and CRM data, building scalable measurement frameworks, and delivering "
    "actionable insights that drive revenue, retention, and Customer Lifetime Value (CLTV). Adept at collaborating "
    "with media, strategy, product, and MarTech teams to launch CDP initiatives, enforce taxonomy/governance standards, "
    "and enable audience-first targeting strategies."
)

# Key Strengths
doc.add_heading('KEY STRENGTHS', level=2)
doc.add_paragraph(
    "• Cross-Channel Media & CRM Measurement (SEM, SEO, Social, OLV, Display)\n"
    "• Attribution & Marketing Science Modeling (MMM, MTA, Causal Inference)\n"
    "• Customer Segmentation (RFM, CLTV, Propensity Scoring)\n"
    "• Campaign Experimentation (A/B & Multivariate Testing, Creative Optimization)\n"
    "• KPI Dashboards & Data Visualization (Tableau, Power BI)\n"
    "• Taxonomy & Data Governance for Marketing Analytics\n"
    "• Stakeholder Communication & Data Storytelling"
)

# Tools & Technologies
doc.add_heading('TOOLS & TECHNOLOGIES', level=2)
doc.add_paragraph(
    "Data & Marketing Science: MMM, MTA, Attribution, CRM Analysis, Audience Segmentation, CDP Strategy\n"
    "Platforms: Google Marketing Platform, Meta Business Suite, TikTok Ads, Salesforce Marketing Cloud, Adobe Analytics\n"
    "Databases & Cloud: Google BigQuery, Azure Databricks, Apache Spark, Hadoop\n"
    "Visualization & BI: Tableau, Power BI\n"
    "Programming: Python, R, SQL, Shell"
)

# Professional Experience
doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)

# TD Role
doc.add_heading('Manager, Data Science & Analytics', level=3)
doc.add_paragraph('Toronto-Dominion Canada Trust | Nov 2021 – Jun 2024')
doc.add_paragraph(
    "• Enhanced digital campaign measurement for personalized marketing channels (145M+ annual emails) by linking conversions "
    "to attributed revenue from cross-sell and upsell campaigns.\n"
    "• Partnered with product marketing teams to measure the effect of paid media & CRM touchpoints on onboarding, adoption, "
    "and cross-sell behavior.\n"
    "• Reduced media waste by +15% through optimal frequency capping and targeted audience analysis across SEM, Social, Display, "
    "and OLV campaigns using Python automation.\n"
    "• Built Power BI and Tableau dashboards integrating first-party CRM, agency, and product engagement data to monitor KPIs "
    "and connect marketing to downstream product behavior.\n"
    "• Measured $200M in digital ad-driven sales using multi-touch attribution (MTA) and statistical modeling in Python/SQL.\n"
    "• Collaborated with MarTech teams on CDP data integration, ensuring taxonomy governance, tagging consistency, and clean "
    "cross-channel measurement.\n"
    "• Designed creative and audience-focused A/B tests to optimize messaging, visual design, and frequency strategies."
)

# Epsilon Role
doc.add_heading('Team Lead – Media Analytics', level=3)
doc.add_paragraph('Epsilon-Publicis Groupe | Jan 2019 – Dec 2020')
doc.add_paragraph(
    "• Led a team of 5 analysts delivering actionable insights across paid search, social, display, and video channels.\n"
    "• Optimized paid social campaigns ($0.65M quarterly spend) improving CTR by +5.6% via creative and targeting adjustments.\n"
    "• Designed and deployed MMM & attribution models for D2C clients, delivering optimal budget allocation and incremental impact reporting.\n"
    "• Developed media measurement frameworks, KPI definitions, and optimization strategies to guide client marketing investments.\n"
    "• Integrated first-party CRM data with media exposure datasets to evaluate campaign ROI and CLTV uplift.\n"
    "• Delivered data-driven storytelling to executives, influencing targeting and media mix strategies."
)

# Hearst Role
doc.add_heading('Manager – Data Science & Media Insights', level=3)
doc.add_paragraph('Hearst Media | 2016 – 2018')
doc.add_paragraph(
    "• Managed $29.4M in paid media investments, producing quarterly forecasts and optimization recommendations using ML models.\n"
    "• Launched audience segmentation from a Customer Data Platform (CDP), increasing LTV accuracy by +15% and improving retention strategies.\n"
    "• Designed incremental lift tests across paid display, paid social, and retargeting, generating measurable gains in site visits and brand recall.\n"
    "• Implemented Adobe Analytics for cross-platform behavior tracking, integrated with CRM for unified customer journey measurement.\n"
    "• Delivered creative-focused A/B testing insights that lifted conversion rates by +15% across key campaigns."
)

# Nielsen Role
doc.add_heading('Senior Data Analyst', level=3)
doc.add_paragraph('Nielsen Marketing Services | 2015 – 2016')
doc.add_paragraph(
    "• Optimized marketing budget allocation across digital (search, display, social) and traditional channels by +4.5% using media saturation curves.\n"
    "• Managed paid media analytics and reporting frameworks, ensuring clean, accurate measurement.\n"
    "• Ran digital media attribution (MTA) for Danone’s CPG brands, integrating store-level, market-level, and CRM sales data.\n"
    "• Contributed to MMM projects for Anheuser-Busch InBev, quantifying pricing, trade, and distribution impacts on performance."
)

# Education & Certs
doc.add_heading('EDUCATION & CERTIFICATIONS', level=2)
doc.add_paragraph(
    "• Master of Engineering, Industrial (Statistics Minor) – Texas A&M University, USA\n"
    "• Bachelor of Engineering, Mechanical – Visvesvaraya Technological University, India\n"
    "• Azure Machine Learning Associate | Google Analytics Certified | Meta Certified Marketing Science Professional | "
    "Google Cloud Data Analytics Professional | Tableau Data Analyst"
)

# Publications
doc.add_heading('PUBLICATIONS', level=2)
doc.add_paragraph(
    "• Machine Learning Beginner’s Guide to Learning All the Fundamental Concepts: 30 Days Challenge\n"
    "• Structured Query Language: Guide to Learning All the Fundamental Concepts: 30 Days Challenge"
)

# Save the file
output_path = 'Heath_Resume_Marketing_Science_Director.docx'
doc.save(output_path)

output_path